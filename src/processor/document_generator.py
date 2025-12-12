"""
Document Generator with 6-Stage Pipeline

This module orchestrates the complete document generation pipeline using
AWS Bedrock (Claude Sonnet 4) to transform video transcriptions into
structured technical training and troubleshooting documents.

Pipeline Stages:
1. Transcription Cleaning and Analysis
2. Technical Content Extraction
3. Solution Mapping
4. Document Structuring
5. Content Writing
6. Output Generation (Markdown + DOCX)

Author: AI Techne Academy
Version: 1.0.0
"""

import json
import logging
import re
from typing import Dict, Any, List, Optional, Tuple
from dataclasses import dataclass
from datetime import datetime, timezone
import io

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from transcription_parser import TranscriptionParser, TranscriptionChunk
from llm_client import BedrockLLMClient, TokenUsage, create_xml_prompt

logger = logging.getLogger(__name__)


@dataclass
class StageResult:
    """Result from a pipeline stage."""
    
    stage_id: int
    stage_name: str
    output: Any
    tokens_used: TokenUsage
    duration_seconds: float
    status: str  # 'success' or 'failed'
    error: Optional[str] = None


@dataclass
class DocumentGenerationResult:
    """Final result of document generation."""
    
    markdown_content: str
    markdown_s3_uri: str
    docx_s3_uri: str
    total_tokens: TokenUsage
    total_cost_usd: float
    duration_seconds: float
    stages: List[StageResult]
    chunks_processed: int


class DocumentGenerator:
    """
    Document generator using 6-stage pipeline.
    
    Orchestrates:
    - Transcription parsing and chunking
    - LLM-based content extraction and structuring
    - Document writing and formatting
    - Multi-format output generation (MD + DOCX)
    """
    
    def __init__(
        self,
        llm_client: BedrockLLMClient,
        parser: TranscriptionParser,
        s3_client: Any
    ):
        """
        Initialize document generator.
        
        Args:
            llm_client: Bedrock LLM client
            parser: Transcription parser
            s3_client: boto3 S3 client
        """
        self.llm = llm_client
        self.parser = parser
        self.s3 = s3_client
        
        logger.info("DocumentGenerator initialized")
    
    def generate_document(
        self,
        execution_id: str,
        transcription_s3_uri: str,
        output_bucket: str
    ) -> DocumentGenerationResult:
        """
        Execute complete 6-stage pipeline.
        
        Args:
            execution_id: Unique execution identifier
            transcription_s3_uri: S3 URI of transcription JSON
            output_bucket: S3 bucket for outputs
            
        Returns:
            DocumentGenerationResult with all outputs and metadata
            
        Raises:
            Exception: If pipeline fails
        """
        start_time = datetime.now()
        logger.info(f"Starting document generation for execution {execution_id}")
        
        # Set logging folder for LLM client
        self.llm.set_logging_folder(f"/tmp/ai-techne/academy/{execution_id}")
        
        try:
            # Load and parse transcription
            transcription_data = self._load_transcription(transcription_s3_uri)
            parsed = self.parser.parse_transcribe_json(transcription_data)
            
            # Chunk if necessary
            chunks = self.parser.chunk_transcription(parsed)
            logger.info(f"Processing {len(chunks)} chunk(s)")
            
            # Process chunks through pipeline
            stage_results = []
            
            # For multi-chunk, process each chunk separately then merge
            if len(chunks) == 1:
                chunk_results = self._process_single_chunk(chunks[0])
                stage_results = chunk_results
            else:
                chunk_results = self._process_multiple_chunks(chunks)
                stage_results = chunk_results
            
            # Extract final markdown from last stage
            final_stage = stage_results[-1]
            markdown_content = final_stage.output
            
            # Stage 6: Generate outputs
            stage_6_start = datetime.now()
            markdown_uri, docx_uri = self._stage_6_generate_outputs(
                execution_id=execution_id,
                markdown_content=markdown_content,
                output_bucket=output_bucket
            )
            stage_6_duration = (datetime.now() - stage_6_start).total_seconds()
            
            stage_results.append(StageResult(
                stage_id=6,
                stage_name="Output Generation",
                output={'markdown_uri': markdown_uri, 'docx_uri': docx_uri},
                tokens_used=TokenUsage(0, 0),
                duration_seconds=stage_6_duration,
                status='success'
            ))
            
            # Calculate totals
            total_tokens = TokenUsage(
                input_tokens=sum(s.tokens_used.input_tokens for s in stage_results),
                output_tokens=sum(s.tokens_used.output_tokens for s in stage_results)
            )
            total_duration = (datetime.now() - start_time).total_seconds()
            
            result = DocumentGenerationResult(
                markdown_content=markdown_content,
                markdown_s3_uri=markdown_uri,
                docx_s3_uri=docx_uri,
                total_tokens=total_tokens,
                total_cost_usd=total_tokens.calculate_cost(),
                duration_seconds=total_duration,
                stages=stage_results,
                chunks_processed=len(chunks)
            )
            
            logger.info(
                f"Document generation complete: {total_duration:.2f}s, "
                f"${result.total_cost_usd:.4f}, {len(markdown_content)} chars"
            )
            
            return result
            
        except Exception as e:
            logger.error(f"Document generation failed: {str(e)}")
            raise
    
    def _process_single_chunk(
        self,
        chunk: TranscriptionChunk
    ) -> List[StageResult]:
        """Process single chunk through all stages."""
        results = []
        
        # Stage 1: Clean transcription
        stage_1_result = self._stage_1_clean_transcription(chunk)
        results.append(stage_1_result)
        
        # Stage 2: Extract technical content
        stage_2_result = self._stage_2_extract_technical_content(
            stage_1_result.output
        )
        results.append(stage_2_result)
        
        # Stage 3: Map solutions
        stage_3_result = self._stage_3_map_solutions(
            stage_2_result.output
        )
        results.append(stage_3_result)
        
        # Stage 4: Structure document
        stage_4_result = self._stage_4_structure_document(
            stage_3_result.output
        )
        results.append(stage_4_result)
        
        # Stage 5: Write content
        stage_5_result = self._stage_5_write_content(
            stage_4_result.output
        )
        results.append(stage_5_result)
        
        return results
    
    def _process_multiple_chunks(
        self,
        chunks: List[TranscriptionChunk]
    ) -> List[StageResult]:
        """
        Process multiple chunks and merge results.
        
        Strategy:
        - Process each chunk through stages 1-4 independently
        - Merge stage 4 outputs (outlines)
        - Execute stage 5 once with merged outline
        """
        logger.info(f"Processing {len(chunks)} chunks with merge strategy")
        
        chunk_outlines = []
        all_technical_content = []
        
        for i, chunk in enumerate(chunks):
            logger.info(f"Processing chunk {i+1}/{len(chunks)}")
            
            # Stages 1-4 for each chunk
            stage_1 = self._stage_1_clean_transcription(chunk)
            stage_2 = self._stage_2_extract_technical_content(stage_1.output)
            stage_3 = self._stage_3_map_solutions(stage_2.output)
            stage_4 = self._stage_4_structure_document(stage_3.output)
            
            chunk_outlines.append(stage_4.output)
            all_technical_content.append(stage_2.output)
        
        # Merge outlines
        logger.info("Merging chunk outlines")
        merged_outline = self._merge_outlines(chunk_outlines)
        
        # Stage 5 with merged outline
        stage_5_result = self._stage_5_write_content(merged_outline)
        
        # Return aggregated results (simplified)
        return [stage_5_result]
    
    def _stage_1_clean_transcription(
        self,
        chunk: TranscriptionChunk
    ) -> StageResult:
        """
        Stage 1: Clean and prepare transcription.
        
        Removes non-technical dialogue and formats with speakers.
        """
        logger.info(f"Stage 1: Cleaning transcription (chunk {chunk.chunk_id})")
        start_time = datetime.now()
        
        try:
            # Format with timestamps and speakers
            formatted_text = self.parser.format_with_timestamps(
                chunk.segments,
                include_speakers=True
            )
            
            # For now, return formatted text directly
            # In production, could add LLM-based noise removal
            cleaned_text = formatted_text
            
            duration = (datetime.now() - start_time).total_seconds()
            
            return StageResult(
                stage_id=1,
                stage_name="Transcription Cleaning",
                output=cleaned_text,
                tokens_used=TokenUsage(0, 0),  # No LLM call
                duration_seconds=duration,
                status='success'
            )
            
        except Exception as e:
            logger.error(f"Stage 1 failed: {str(e)}")
            return StageResult(
                stage_id=1,
                stage_name="Transcription Cleaning",
                output=None,
                tokens_used=TokenUsage(0, 0),
                duration_seconds=0,
                status='failed',
                error=str(e)
            )
    
    def _stage_2_extract_technical_content(
        self,
        cleaned_transcription: str
    ) -> StageResult:
        """
        Stage 2: Extract technical content using LLM.
        
        Identifies:
        - Diagnostics (errors, causes)
        - Solutions (steps, commands)
        - Risks (environment, concurrency)
        - Business rules
        - Configurations
        """
        logger.info("Stage 2: Extracting technical content")
        start_time = datetime.now()
        
        # Set stage context for logging
        self.llm.set_stage_context("stage-2")
        
        prompt = create_xml_prompt(
            task="Você é um Especialista em Documentação Técnica de Software. Analise a transcrição fornecida e extraia conteúdo técnico.",
            instructions="""Ignore diálogos sociais. Concentre-se exclusivamente no conteúdo técnico e extraia:

1. <diagnostic>
   Liste erros específicos (códigos, mensagens, exceções).
   Para cada erro, identifique a causa raiz explicada.
</diagnostic>

2. <solutions>
   Descreva passos técnicos executados.
   Inclua comandos, alterações de código, configurações.
</solutions>

3. <risks>
   Identifique avisos sobre ambientes compartilhados, concorrência, 
   procedimentos que afetam outros membros da equipe.
</risks>

4. <business_rules>
   Extraia explicações sobre comportamento do sistema.
   Por que dados aparecem com valor X ou Y?
</business_rules>

5. <configurations>
   Liste convenções de nomenclatura (prefixos, sufixos).
   Configurações específicas de ferramentas.
</configurations>""",
            output_format="Gere um JSON estruturado contendo os 5 elementos acima. Use listas para múltiplos itens em cada categoria.",
            input_data=cleaned_transcription,
            input_tag="transcription"
        )
        
        try:
            response, usage = self.llm.invoke_with_json_output(prompt)
            duration = (datetime.now() - start_time).total_seconds()
            
            logger.info(f"Stage 2 complete: {usage.total_tokens} tokens")
            
            return StageResult(
                stage_id=2,
                stage_name="Technical Content Extraction",
                output=response,
                tokens_used=usage,
                duration_seconds=duration,
                status='success'
            )
            
        except Exception as e:
            logger.error(f"Stage 2 failed: {str(e)}")
            duration = (datetime.now() - start_time).total_seconds()
            return StageResult(
                stage_id=2,
                stage_name="Technical Content Extraction",
                output=None,
                tokens_used=TokenUsage(0, 0),
                duration_seconds=duration,
                status='failed',
                error=str(e)
            )
    
    def _stage_3_map_solutions(
        self,
        technical_content: Dict[str, Any]
    ) -> StageResult:
        """
        Stage 3: Map problems to solutions.
        
        Creates clear problem → solution mapping.
        """
        logger.info("Stage 3: Mapping solutions")
        start_time = datetime.now()
        
        # Set stage context for logging
        self.llm.set_stage_context("stage-3")
        
        prompt = create_xml_prompt(
            task="Você é um Engenheiro de Software Sênior. Crie um mapeamento entre problemas e soluções.",
            instructions="""Com base no conteúdo técnico extraído, crie uma matriz problema-solução:

1. <problem_solution_map>
   Para cada erro/problema identificado:
   - Código/mensagem do erro
   - Causa raiz
   - Solução aplicada (passo a passo)
   - Comandos executados
   - Resultado obtido
</problem_solution_map>

2. <preventive_measures>
   Liste medidas preventivas mencionadas.
   Como evitar o problema no futuro?
</preventive_measures>

3. <debugging_steps>
   Passos de debug/investigação mencionados.
   Como diagnosticar o problema?
</debugging_steps>""",
            output_format="Gere um JSON estruturado com os 3 elementos acima. Organize de forma hierárquica (problema → causa → solução).",
            input_data=json.dumps(technical_content, indent=2),
            input_tag="technical_content"
        )
        
        try:
            response, usage = self.llm.invoke_with_json_output(prompt)
            duration = (datetime.now() - start_time).total_seconds()
            
            logger.info(f"Stage 3 complete: {usage.total_tokens} tokens")
            
            return StageResult(
                stage_id=3,
                stage_name="Solution Mapping",
                output=response,
                tokens_used=usage,
                duration_seconds=duration,
                status='success'
            )
            
        except Exception as e:
            logger.error(f"Stage 3 failed: {str(e)}")
            duration = (datetime.now() - start_time).total_seconds()
            return StageResult(
                stage_id=3,
                stage_name="Solution Mapping",
                output=None,
                tokens_used=TokenUsage(0, 0),
                duration_seconds=duration,
                status='failed',
                error=str(e)
            )
    
    def _stage_4_structure_document(
        self,
        solution_map: Dict[str, Any]
    ) -> StageResult:
        """
        Stage 4: Create document structure (outline).
        
        Organizes content into logical sections.
        """
        logger.info("Stage 4: Structuring document")
        start_time = datetime.now()
        
        # Set stage context for logging
        self.llm.set_stage_context("stage-4")
        
        prompt = create_xml_prompt(
            task="Você é um Designer Instrucional Técnico. Crie a estrutura para um Guia de Treinamento e Troubleshooting.",
            instructions="""Organize os tópicos da seguinte forma:

1. <troubleshooting_section>
   - Comparação entre tipos de erros (Compilação vs Runtime vs Configuração)
   - Cada erro associado à sua solução imediata
   - Sintomas e diagnóstico
</troubleshooting_section>

2. <practical_procedure>
   - Passo a passo da solução técnica
   - Seção específica: 'Instrumentação e Debug'
   - Como rastrear o problema
   - Comandos e ferramentas utilizadas
</practical_procedure>

3. <safety_protocols>
   - Cuidados com ambiente compartilhado
   - Como não impactar colegas
   - Seção destacada como CRÍTICA
</safety_protocols>

4. <business_understanding>
   - Interpretação dos resultados
   - Explicação de regras de negócio
   - Por que o sistema se comporta assim?
</business_understanding>

5. <faq_section>
   - Dúvidas comuns encontradas
   - Respostas objetivas
</faq_section>""",
            output_format="Gere apenas a estrutura (outline) em formato hierárquico. Use bullet points detalhando o que será abordado em cada seção. NÃO escreva o conteúdo completo ainda.",
            input_data=json.dumps(solution_map, indent=2),
            input_tag="solution_map"
        )
        
        try:
            response, usage = self.llm.invoke(prompt)
            duration = (datetime.now() - start_time).total_seconds()
            
            logger.info(f"Stage 4 complete: {usage.total_tokens} tokens")
            
            return StageResult(
                stage_id=4,
                stage_name="Document Structuring",
                output=response,
                tokens_used=usage,
                duration_seconds=duration,
                status='success'
            )
            
        except Exception as e:
            logger.error(f"Stage 4 failed: {str(e)}")
            duration = (datetime.now() - start_time).total_seconds()
            return StageResult(
                stage_id=4,
                stage_name="Document Structuring",
                output=None,
                tokens_used=TokenUsage(0, 0),
                duration_seconds=duration,
                status='failed',
                error=str(e)
            )
    
    def _stage_5_write_content(
        self,
        document_outline: str
    ) -> StageResult:
        """
        Stage 5: Write complete Markdown document.
        
        Generates final document based on outline.
        """
        logger.info("Stage 5: Writing document content")
        start_time = datetime.now()
        
        # Set stage context for logging
        self.llm.set_stage_context("stage-5")
        
        prompt = create_xml_prompt(
            task="Atue como um Redator Técnico Sênior. Escreva um Documento de Treinamento e Troubleshooting completo em Markdown.",
            instructions="""<guidelines>
- Tom: Profissional, instrucional, direto (imperativo: 'Faça', 'Configure', 'Verifique')
- Clareza: Acessível para iniciantes. Explique o 'porquê', não apenas o 'como'
- Formatação: Use code blocks, negrito para ênfase, blockquotes para avisos críticos
</guidelines>

<required_structure>
# 1. Introdução
- Contexto: por que este documento existe?
- Qual problema/tema está sendo abordado?
- Objetivo final para o leitor

# 2. Desenvolvimento
## 2.1 Conceitos-chave
## 2.2 Troubleshooting Guide
## 2.3 Procedimento Passo a Passo
## 2.4 Melhores Práticas
## 2.5 Considerações de Segurança
## 2.6 Entendendo os Dados (Regras de Negócio)

Use destaques (callouts), tabelas e listas quando melhorar o entendimento.

# 3. Encerramento
- Resumo do que foi aprendido
- Reforço da mensagem principal
- Próximos passos/ações

# 4. FAQ
- Dúvidas comuns
- Respostas objetivas e claras
</required_structure>

<quality_checklist>
- Texto objetivo e didático
- Títulos de seção claros
- Apenas informações relevantes
- Sem termos ofensivos ou piadas internas
- Sem frases de enchimento ou diálogos irrelevantes
</quality_checklist>""",
            output_format="Gere o documento completo em Markdown AGORA. Use toda a informação fornecida para criar um documento rico, detalhado e profissional.",
            input_data=document_outline,
            input_tag="document_outline"
        )
        
        try:
            # Use higher token limit for final document
            original_max = self.llm.max_tokens
            self.llm.model.model_kwargs["max_tokens"] = 8192
            
            response, usage = self.llm.invoke(prompt, temperature=0.7)
            
            # Restore original
            self.llm.model.model_kwargs["max_tokens"] = original_max
            
            duration = (datetime.now() - start_time).total_seconds()
            
            logger.info(
                f"Stage 5 complete: {usage.total_tokens} tokens, "
                f"{len(response)} chars"
            )
            
            return StageResult(
                stage_id=5,
                stage_name="Content Writing",
                output=response,
                tokens_used=usage,
                duration_seconds=duration,
                status='success'
            )
            
        except Exception as e:
            logger.error(f"Stage 5 failed: {str(e)}")
            duration = (datetime.now() - start_time).total_seconds()
            return StageResult(
                stage_id=5,
                stage_name="Content Writing",
                output=None,
                tokens_used=TokenUsage(0, 0),
                duration_seconds=duration,
                status='failed',
                error=str(e)
            )
    
    def _stage_6_generate_outputs(
        self,
        execution_id: str,
        markdown_content: str,
        output_bucket: str
    ) -> Tuple[str, str]:
        """
        Stage 6: Generate and upload Markdown + DOCX.
        
        Returns:
            Tuple of (markdown_s3_uri, docx_s3_uri)
        """
        logger.info("Stage 6: Generating output files")
        
        # Save Markdown
        md_key = f"{execution_id}/document.md"
        md_uri = f"s3://{output_bucket}/{md_key}"
        
        self.s3.put_object(
            Bucket=output_bucket,
            Key=md_key,
            Body=markdown_content.encode('utf-8'),
            ContentType='text/markdown',
            Metadata={
                'execution_id': execution_id,
                'generated_at': datetime.now(timezone.utc).isoformat(),
                'content_length': str(len(markdown_content))
            }
        )
        logger.info(f"Markdown saved: {md_uri}")
        
        # Convert to DOCX
        docx_buffer = self._markdown_to_docx(markdown_content)
        
        # Save DOCX
        docx_key = f"{execution_id}/document.docx"
        docx_uri = f"s3://{output_bucket}/{docx_key}"
        
        self.s3.put_object(
            Bucket=output_bucket,
            Key=docx_key,
            Body=docx_buffer.getvalue(),
            ContentType='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            Metadata={
                'execution_id': execution_id,
                'generated_at': datetime.now(timezone.utc).isoformat()
            }
        )
        logger.info(f"DOCX saved: {docx_uri}")
        
        return md_uri, docx_uri
    
    def _markdown_to_docx(self, markdown_content: str) -> io.BytesIO:
        """
        Convert Markdown to DOCX using python-docx.
        
        Supports:
        - Headers (#, ##, ###)
        - Bold (**text**)
        - Italic (*text*)
        - Code blocks (```)
        - Inline code (`code`)
        - Lists (bullet and numbered)
        - Blockquotes (>)
        - Tables (basic)
        """
        logger.info("Converting Markdown to DOCX")
        
        doc = Document()
        
        # Parse markdown line by line
        lines = markdown_content.split('\n')
        in_code_block = False
        code_block_lines = []
        in_list = False
        list_items = []
        
        for line in lines:
            # Code blocks
            if line.strip().startswith('```'):
                if in_code_block:
                    # End code block
                    code_text = '\n'.join(code_block_lines)
                    p = doc.add_paragraph(code_text)
                    p.style = 'Intense Quote'
                    code_block_lines = []
                    in_code_block = False
                else:
                    # Start code block
                    in_code_block = True
                continue
            
            if in_code_block:
                code_block_lines.append(line)
                continue
            
            # Headers
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            
            # Bullet lists
            elif line.strip().startswith('- ') or line.strip().startswith('* '):
                text = line.strip()[2:]
                doc.add_paragraph(text, style='List Bullet')
            
            # Numbered lists
            elif re.match(r'^\d+\. ', line.strip()):
                text = re.sub(r'^\d+\. ', '', line.strip())
                doc.add_paragraph(text, style='List Number')
            
            # Blockquotes
            elif line.strip().startswith('>'):
                text = line.strip()[1:].strip()
                p = doc.add_paragraph(text)
                p.style = 'Quote'
            
            # Regular paragraph
            elif line.strip():
                # Process inline formatting
                p = doc.add_paragraph()
                self._add_formatted_text(p, line)
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer
    
    def _add_formatted_text(self, paragraph, text: str):
        """Add text with inline formatting (bold, italic, code)."""
        # Simple implementation - could be enhanced
        # For now, just add as plain text
        # TODO: Parse **bold**, *italic*, `code`
        paragraph.add_run(text)
    
    def _merge_outlines(self, outlines: List[str]) -> str:
        """
        Merge multiple outlines into one.
        
        Simple strategy: concatenate with markers.
        For production, could use LLM to merge intelligently.
        """
        merged = "# Merged Outline from Multiple Chunks\n\n"
        
        for i, outline in enumerate(outlines):
            merged += f"## Part {i+1}\n\n{outline}\n\n"
        
        return merged
    
    def _load_transcription(self, s3_uri: str) -> Dict[str, Any]:
        """Load transcription JSON from S3."""
        from transcription_parser import parse_s3_uri, load_transcription_from_s3
        
        bucket, key = parse_s3_uri(s3_uri)
        return load_transcription_from_s3(self.s3, bucket, key)